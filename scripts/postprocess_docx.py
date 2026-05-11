"""Постобработка thesis_vkr.docx после pandoc.

Доработки под методичку РАНХиГС 2023:
1. Нумерация страниц по центру верхнего поля (со 2-й страницы — на титуле
   номера не будет; titlePg = особый колонтитул для 1-й страницы).
2. Разрыв страницы перед каждой Heading 1, кроме первой («Введение»).
3. Подписи «Таблица N.N — ...» и «Рисунок N.N — ...» по центру, без
   абзацного отступа, без курсива.
4. Маркер «–» (en-dash) для всех ненумерованных списков (§2.2.6).
5. «Введение», «Заключение», «Список литературы», «Оглавление» —
   по центру, без абзацного отступа.
"""

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


CAPTION_RE = re.compile(r"^(Таблица|Рисунок)\s+\d+(\.\d+)*\s*[—–-]")


def add_page_numbers(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in list(p.runs):
        run.text = ""
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    sectPr = section._sectPr
    titlePg = sectPr.find(qn("w:titlePg"))
    if titlePg is None:
        titlePg = OxmlElement("w:titlePg")
        sectPr.append(titlePg)


def clear_first_line_indent(para):
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    for attr in ("w:firstLine", "w:left", "w:hanging"):
        if qn(attr) in ind.attrib:
            del ind.attrib[qn(attr)]


def add_page_break_before(para):
    pPr = para._p.get_or_add_pPr()
    pbb = pPr.find(qn("w:pageBreakBefore"))
    if pbb is None:
        pbb = OxmlElement("w:pageBreakBefore")
        pPr.append(pbb)


def add_breaks_before_headings(doc):
    first_h1_seen = False
    for para in doc.paragraphs:
        if para.style and para.style.name == "Heading 1":
            if not first_h1_seen:
                first_h1_seen = True
                continue
            add_page_break_before(para)


def center_captions(doc):
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if CAPTION_RE.match(text):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            clear_first_line_indent(para)
            for run in para.runs:
                run.italic = False


def fix_special_headings(doc):
    targets = {"введение", "заключение", "список литературы", "оглавление"}
    for para in doc.paragraphs:
        if not para.style or not para.style.name.startswith("Heading"):
            continue
        if para.text.strip().lower() in targets:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            clear_first_line_indent(para)


def restructure_numbered_equations(doc):
    """Помещает пронумерованные display-формулы в безрамочную таблицу 2 столбца:
    левая ячейка (широкая) — уравнение по центру,
    правая ячейка (узкая) — номер вида (N.N) по правому краю.
    """
    from copy import deepcopy

    MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    NUM_RE = re.compile(r"\((\d+\.\d+)\)")

    paras_to_replace = []

    for para in doc.paragraphs:
        omath_para = para._p.find(".//{%s}oMathPara" % MATH_NS)
        if omath_para is None:
            continue
        omath = omath_para.find("{%s}oMath" % MATH_NS)
        if omath is None:
            continue

        # Собираем текст из хвостовых m:r-элементов, ищем (N.N)
        children = list(omath)
        tail_text = ""
        tail_elems = []

        for child in reversed(children):
            if child.tag != "{%s}r" % MATH_NS:
                break
            t_el = child.find("{%s}t" % MATH_NS)
            text = (t_el.text or "") if t_el is not None else ""
            tail_elems.insert(0, child)
            tail_text = text + tail_text
            if NUM_RE.search(tail_text):
                break

        m = NUM_RE.search(tail_text)
        if not m:
            continue

        num_label = f"({m.group(1)})"
        paras_to_replace.append((para, omath, tail_elems, num_label))

    # страничная ширина минус поля: A4 210мм − 35мм − 10мм = 165мм = 9360 twips
    PAGE_W = 9360
    COL1 = int(PAGE_W * 0.87)
    COL2 = PAGE_W - COL1

    for para, omath, tail_elems, num_label in paras_to_replace:
        # Удаляем qquad + номер из OMML
        for elem in tail_elems:
            try:
                omath.remove(elem)
            except ValueError:
                pass

        # Строим таблицу 1×2 без рамок
        tbl = OxmlElement("w:tbl")

        tblPr = OxmlElement("w:tblPr")
        tblW_el = OxmlElement("w:tblW")
        tblW_el.set(qn("w:w"), str(PAGE_W))
        tblW_el.set(qn("w:type"), "dxa")
        tblPr.append(tblW_el)

        tblBorders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "none")
            b.set(qn("w:sz"), "0")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "auto")
            tblBorders.append(b)
        tblPr.append(tblBorders)

        tblCellMar = OxmlElement("w:tblCellMar")
        for side in ("top", "left", "bottom", "right"):
            m_el = OxmlElement(f"w:{side}")
            m_el.set(qn("w:w"), "0")
            m_el.set(qn("w:type"), "dxa")
            tblCellMar.append(m_el)
        tblPr.append(tblCellMar)
        tbl.append(tblPr)

        tblGrid = OxmlElement("w:tblGrid")
        for w in (COL1, COL2):
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(w))
            tblGrid.append(gc)
        tbl.append(tblGrid)

        tr = OxmlElement("w:tr")

        # --- Ячейка 1: формула по центру ---
        tc1 = OxmlElement("w:tc")
        tc1Pr = OxmlElement("w:tcPr")
        tc1W = OxmlElement("w:tcW")
        tc1W.set(qn("w:w"), str(COL1))
        tc1W.set(qn("w:type"), "dxa")
        tc1Pr.append(tc1W)
        tc1.append(tc1Pr)

        fp = deepcopy(para._p)
        pPr_fp = fp.find(qn("w:pPr"))
        if pPr_fp is None:
            pPr_fp = OxmlElement("w:pPr")
            fp.insert(0, pPr_fp)
        jc_fp = pPr_fp.find(qn("w:jc"))
        if jc_fp is None:
            jc_fp = OxmlElement("w:jc")
            pPr_fp.append(jc_fp)
        jc_fp.set(qn("w:val"), "center")
        tc1.append(fp)
        tr.append(tc1)

        # --- Ячейка 2: номер по правому краю ---
        tc2 = OxmlElement("w:tc")
        tc2Pr = OxmlElement("w:tcPr")
        tc2W = OxmlElement("w:tcW")
        tc2W.set(qn("w:w"), str(COL2))
        tc2W.set(qn("w:type"), "dxa")
        tc2Pr.append(tc2W)
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), "center")
        tc2Pr.append(vAlign)
        tc2.append(tc2Pr)

        num_p = OxmlElement("w:p")
        num_pPr = OxmlElement("w:pPr")
        num_jc = OxmlElement("w:jc")
        num_jc.set(qn("w:val"), "right")
        num_pPr.append(num_jc)
        num_ind = OxmlElement("w:ind")
        num_ind.set(qn("w:firstLine"), "0")
        num_ind.set(qn("w:left"), "0")
        num_pPr.append(num_ind)
        num_p.append(num_pPr)
        num_r = OxmlElement("w:r")
        num_t = OxmlElement("w:t")
        num_t.text = num_label
        num_r.append(num_t)
        num_p.append(num_r)
        tc2.append(num_p)
        tr.append(tc2)

        tbl.append(tr)

        # Заменяем параграф таблицей
        para._p.addprevious(tbl)
        para._p.getparent().remove(para._p)
    targets = {"введение", "заключение", "список литературы", "оглавление"}
    for para in doc.paragraphs:
        if not para.style or not para.style.name.startswith("Heading"):
            continue
        if para.text.strip().lower() in targets:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            clear_first_line_indent(para)


def style_tables(doc):
    """Включает все границы и автоматическую подгонку ширины колонок.

    Без autofit pandoc разбивает ширину таблицы поровну между колонками,
    из-за чего узкие колонки с числами «−0,054» переносятся по символу.
    Стиль «все границы» отвечает требованию методички: на скриншоте
    видно, что границы должны быть видны полностью.
    """
    MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    NUM_RE = re.compile(r"^\(\d+\.\d+\)$")

    for table in doc.tables:
        # Пропускаем безрамочные таблицы-обёртки для формул
        if (
            len(table.rows) == 1
            and len(table.columns) == 2
            and NUM_RE.match(table.rows[0].cells[1].text.strip())
            and table.rows[0].cells[0]._tc.find(".//{%s}oMathPara" % MATH_NS) is not None
        ):
            continue

        tblPr = table._tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tblPr)

        for border_tag in ("w:tblBorders",):
            existing = tblPr.find(qn(border_tag))
            if existing is not None:
                tblPr.remove(existing)

        tblBorders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "4")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "000000")
            tblBorders.append(b)
        tblPr.append(tblBorders)

        layout = tblPr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tblPr.append(layout)
        layout.set(qn("w:type"), "autofit")

        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), "0")
        tblW.set(qn("w:type"), "auto")

        # Сбрасываем отступ первой строки в стиле параграфа по умолчанию для таблицы
        tblPrChange = tblPr.find(qn("w:tblPrChange"))
        # Устанавливаем нулевые отступы ячеек по умолчанию на уровне таблицы
        tblCellMar = tblPr.find(qn("w:tblCellMar"))
        if tblCellMar is not None:
            tblPr.remove(tblCellMar)
        tblCellMar = OxmlElement("w:tblCellMar")
        for side in ("top", "left", "bottom", "right"):
            m = OxmlElement(f"w:{side}")
            m.set(qn("w:w"), "36")
            m.set(qn("w:type"), "dxa")
            tblCellMar.append(m)
        tblPr.append(tblCellMar)

        tblGrid = table._tbl.find(qn("w:tblGrid"))
        if tblGrid is not None:
            for gridCol in tblGrid.findall(qn("w:gridCol")):
                gridCol.set(qn("w:w"), "0")

        for row in table.rows:
            trPr = row._tr.find(qn("w:trPr"))
            if trPr is None:
                trPr = OxmlElement("w:trPr")
                row._tr.insert(0, trPr)
            cantSplit = trPr.find(qn("w:cantSplit"))
            if cantSplit is None:
                cantSplit = OxmlElement("w:cantSplit")
                trPr.append(cantSplit)

            for cell in row.cells:
                tcPr = cell._tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    cell._tc.insert(0, tcPr)
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is None:
                    tcW = OxmlElement("w:tcW")
                    tcPr.append(tcW)
                tcW.set(qn("w:w"), "0")
                tcW.set(qn("w:type"), "auto")

                noWrap = tcPr.find(qn("w:noWrap"))
                if noWrap is None:
                    noWrap = OxmlElement("w:noWrap")
                    tcPr.append(noWrap)

                # Уменьшаем отступы в ячейке (по 36 twips = ~0.6 мм с каждой стороны)
                tcMar = tcPr.find(qn("w:tcMar"))
                if tcMar is None:
                    tcMar = OxmlElement("w:tcMar")
                    tcPr.append(tcMar)
                for side in ("top", "left", "bottom", "right"):
                    m = tcMar.find(qn(f"w:{side}"))
                    if m is None:
                        m = OxmlElement(f"w:{side}")
                        tcMar.append(m)
                    m.set(qn("w:w"), "36")
                    m.set(qn("w:type"), "dxa")

                # Шрифт 10 пт и сброс абзацного отступа для параграфов внутри ячейки
                for para in cell.paragraphs:
                    pPr = para._p.get_or_add_pPr()

                    # Убираем отступ первой строки (наследуется из Normal — 1.25 см)
                    ind = pPr.find(qn("w:ind"))
                    if ind is None:
                        ind = OxmlElement("w:ind")
                        pPr.append(ind)
                    for attr in ("w:hanging",):
                        if qn(attr) in ind.attrib:
                            del ind.attrib[qn(attr)]
                    ind.set(qn("w:firstLine"), "0")
                    ind.set(qn("w:left"), "0")

                    rPr = pPr.find(qn("w:rPr"))
                    if rPr is None:
                        rPr = OxmlElement("w:rPr")
                        pPr.append(rPr)
                    for sz_tag in ("w:sz", "w:szCs"):
                        sz_el = rPr.find(qn(sz_tag))
                        if sz_el is None:
                            sz_el = OxmlElement(sz_tag)
                            rPr.append(sz_el)
                        sz_el.set(qn("w:val"), "20")  # 10 pt = 20 half-points
                    for run in para.runs:
                        rPr2 = run._r.find(qn("w:rPr"))
                        if rPr2 is None:
                            rPr2 = OxmlElement("w:rPr")
                            run._r.insert(0, rPr2)
                        for sz_tag in ("w:sz", "w:szCs"):
                            sz_el = rPr2.find(qn(sz_tag))
                            if sz_el is None:
                                sz_el = OxmlElement(sz_tag)
                                rPr2.append(sz_el)
                            sz_el.set(qn("w:val"), "20")


def fix_bullet_markers(docx_path: Path):
    """Меняет маркеры всех ненумерованных списков на тире (–).

    В numbering.xml ищет <w:lvl> внутри <w:abstractNum>, где numFmt = bullet,
    и подменяет lvlText на «–». Также удаляет привязку к шрифту Symbol/Wingdings,
    чтобы тире отображалось обычным шрифтом документа.
    """
    import shutil
    import tempfile

    with zipfile.ZipFile(docx_path) as z:
        if "word/numbering.xml" not in z.namelist():
            return
        with z.open("word/numbering.xml") as f:
            numbering = f.read().decode("utf-8")

    bullet_lvl_re = re.compile(
        r"(<w:lvl\b[^>]*>)(.*?)(</w:lvl>)", re.DOTALL
    )

    def fix_lvl(match):
        opening, body, closing = match.group(1), match.group(2), match.group(3)
        if "<w:numFmt w:val=\"bullet\"/>" not in body:
            return match.group(0)
        body = re.sub(
            r'<w:lvlText\s+w:val="[^"]*"\s*/>',
            '<w:lvlText w:val="\u2013"/>',
            body,
        )
        body = re.sub(
            r"<w:rPr>.*?</w:rPr>",
            "",
            body,
            flags=re.DOTALL,
        )
        return opening + body + closing

    new_numbering = bullet_lvl_re.sub(fix_lvl, numbering)

    if new_numbering == numbering:
        return

    tmp = Path(tempfile.mkdtemp()) / "out.docx"
    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(
        tmp, "w", zipfile.ZIP_DEFLATED
    ) as dst:
        for item in src.namelist():
            if item == "word/numbering.xml":
                dst.writestr(item, new_numbering)
            else:
                dst.writestr(item, src.read(item))
    shutil.move(str(tmp), str(docx_path))


def main():
    if len(sys.argv) < 2:
        raise SystemExit("usage: postprocess_docx.py <path>")
    path = Path(sys.argv[1])
    doc = Document(str(path))
    add_page_numbers(doc)
    add_breaks_before_headings(doc)
    center_captions(doc)
    fix_special_headings(doc)
    restructure_numbered_equations(doc)
    style_tables(doc)
    doc.save(str(path))
    fix_bullet_markers(path)
    print(f"Постобработка применена: {path}")


if __name__ == "__main__":
    main()
